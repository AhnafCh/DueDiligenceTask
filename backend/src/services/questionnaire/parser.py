"""Questionnaire Parser using Gemini Flash."""
import logging
from pathlib import Path
from typing import Optional
import json

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain_core.output_parsers import PydanticOutputParser

import PyPDF2
from docx import Document as DocxDocument

from src.utils.config import get_settings
from .models import ParsedQuestionnaire

logger = logging.getLogger(__name__)

class QuestionnaireParser:
    """
    Parses questionnaire documents (PDF/DOCX) into structured data
    using Gemini 1.5 Flash.
    """

    def __init__(self):
        self.settings = get_settings()
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            google_api_key=self.settings.google_api_key,
            temperature=0.0, # Deterministic output
            convert_system_message_to_human=True
        )
        self.parser = PydanticOutputParser(pydantic_object=ParsedQuestionnaire)

        self.prompt_template = PromptTemplate(
            template="""
            You are an expert data extraction assistant. Your task is to extract a structured questionnaire from the provided raw text.

            Identification Rules:
            1. Identify distinct sections. A section usually has a heading (e.g., "1. General Information", "Section B: Financials").
            2. If there are no explicit sections, create a single "General" section.
            3. Identify questions within those sections. Questions usually start with numbers (1, 1.1, a, b) or bullets, or end with a question mark.
            4. Extract the exact text of the question.
            5. Determine the order of sections and questions.
            6. Infer the expected answer type (text, number, boolean, date, etc.) based on the question content.

            {format_instructions}

            Raw Text:
            {raw_text}
            """,
            input_variables=["raw_text"],
            partial_variables={"format_instructions": self.parser.get_format_instructions()}
        )

    def parse_file(self, file_path: Path) -> ParsedQuestionnaire:
        """
        Main entry point to parse a file.
        """
        ext = file_path.suffix.lower()
        if ext == ".pdf":
            raw_text = self._extract_text_from_pdf(file_path)
        elif ext == ".docx":
            raw_text = self._extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        if not raw_text.strip():
            raise ValueError("Could not extract any text from the file.")

        logger.info(f"Extracted {len(raw_text)} chars from {file_path.name}. Sending to Gemini Flash...")
        return self._extract_structure_with_llm(raw_text)

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        text = []
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text())
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            raise IOError(f"Failed to read PDF file: {e}")
        return "\n".join(text)

    def _extract_text_from_docx(self, file_path: Path) -> str:
        text = []
        try:
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    text.append(" | ".join([cell.text for cell in row.cells]))
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            raise IOError(f"Failed to read DOCX file: {e}")
        return "\n".join(text)

    def _extract_structure_with_llm(self, raw_text: str) -> ParsedQuestionnaire:
        try:
            # We might need to split text if it's too long, but Flash has a large context window (1M tokens).
            # For now, we assume the questionnaire fits in context.
            
            prompt = self.prompt_template.format(raw_text=raw_text)
            response = self.llm.invoke(prompt)
            
            # Parse the response
            parsed_data = self.parser.parse(response.content)
            return parsed_data

        except Exception as e:
            logger.error(f"LLM Parsing failed: {e}")
            raise ValueError(f"Failed to parse structure from text: {e}")

