"""DOCX document parser using python-docx."""
import logging
from pathlib import Path
from typing import List
from docx import Document

from .base import BaseParser, ParsedContent

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Parser for Microsoft Word documents."""
    
    def supports(self, file_extension: str) -> bool:
        """Check if file extension is .docx"""
        return file_extension.lower() == ".docx"
    
    def parse(self, file_path: Path) -> ParsedContent:
        """
        Parse DOCX document and extract text with structure.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ParsedContent with extracted text and metadata
            
        Raises:
            ValueError: If DOCX is invalid
            IOError: If file cannot be read
        """
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs: List[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            table_texts: List[str] = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)
            
            # Combine paragraphs and tables
            all_texts = paragraphs + table_texts
            full_text = "\n".join(all_texts)
            cleaned_text = self.clean_text(full_text)
            
            # DOCX doesn't have traditional pages, treat sections as pages
            # For simplicity, we'll consider the entire document as one "page"
            page_texts = [cleaned_text] if cleaned_text else []
            page_count = 1
            
            # Extract metadata
            metadata = self.extract_metadata(
                page_count=page_count,
                paragraph_count=len(paragraphs),
                table_count=len(doc.tables),
                file_size=file_path.stat().st_size,
                file_name=file_path.name
            )
            
            # Add core properties if available
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata['title'] = core_props.title
                if core_props.author:
                    metadata['author'] = core_props.author
                if core_props.created:
                    metadata['created'] = str(core_props.created)
                if core_props.modified:
                    metadata['modified'] = str(core_props.modified)
            except Exception as e:
                logger.debug(f"Could not extract DOCX metadata: {e}")
            
            return ParsedContent(
                text=cleaned_text,
                page_count=page_count,
                metadata=metadata,
                page_texts=page_texts,
                bounding_boxes=None  # DOCX doesn't have bounding boxes
            )
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise IOError(f"Failed to parse DOCX: {e}")
